import io
import json
import logging
import os

from docx import Document
from fpdf import FPDF

logger = logging.getLogger(__name__)

# CJK font path — NotoSansSC bundled in fonts/ directory
_FONT_DIR = os.path.join(os.path.dirname(__file__), "fonts")
_FONT_REGULAR = os.path.join(_FONT_DIR, "NotoSansSC-Regular.ttf")
_FONT_BOLD = os.path.join(_FONT_DIR, "NotoSansSC-Bold.ttf")
_HAS_CJK_FONT = os.path.isfile(_FONT_REGULAR)


def _normalize_citation(c) -> str:
    """Convert a citation (dict or string) to a display string."""
    if isinstance(c, dict):
        title = c.get("title", "")
        link = c.get("link", "")
        source = c.get("source_type", "")
        return f"{title} - {link} ({source})" if link else f"{title} ({source})"
    return str(c)


def _normalize_citation_md(c) -> str:
    """Convert a citation to Markdown link format."""
    if isinstance(c, dict):
        title = c.get("title", "")
        link = c.get("link", "")
        source = c.get("source_type", "")
        return f"- [{title}]({link}) ({source})" if link else f"- {title} ({source})"
    return f"- {c}"


def _safe_str(val) -> str:
    """Convert any value to a display string — handles objects, lists, None."""
    if val is None:
        return ""
    if isinstance(val, str):
        return val
    if isinstance(val, dict):
        return val.get("overall_assessment") or val.get("summary") or json.dumps(val, ensure_ascii=False)
    return str(val)


def generate_markdown_report(report: dict) -> str:
    """Generate a Markdown report from the decision output."""
    citations = report.get("citations", [])
    if not isinstance(citations, list):
        citations = []
    citation_lines = "\n".join(_normalize_citation_md(c) for c in citations)
    risks_list = report.get("major_risks", [])
    if not isinstance(risks_list, list):
        risks_list = [str(risks_list)] if risks_list else []
    risks = "\n".join(f"- {_safe_str(r)}" for r in risks_list)
    opps_list = report.get("major_opportunities", [])
    if not isinstance(opps_list, list):
        opps_list = [str(opps_list)] if opps_list else []
    opportunities = "\n".join(f"- {_safe_str(o)}" for o in opps_list)

    return f"""# Drug Target Assessment Report

## Target: {report.get('target', '')}
## Indication: {report.get('indication', '')}

---

## Literature Summary
{_safe_str(report.get('literature_summary', ''))}

## Clinical Trials Summary
{_safe_str(report.get('clinical_trials_summary', ''))}

## Competition Summary
{_safe_str(report.get('competition_summary', ''))}

## Major Risks
{risks}

## Major Opportunities
{opportunities}

## Recommendation: {report.get('recommendation', '')}

**Reasoning:** {report.get('reasoning', '')}

**Uncertainty:** {report.get('uncertainty', '')}

## Citations
{citation_lines}
"""


def generate_word_report(report: dict) -> bytes:
    """Generate a Word document from the decision output."""
    doc = Document()
    doc.add_heading("Drug Target Assessment Report", level=0)

    doc.add_heading(f"Target: {report.get('target', '')}", level=1)
    doc.add_heading(f"Indication: {report.get('indication', '')}", level=1)

    doc.add_heading("Literature Summary", level=2)
    doc.add_paragraph(_safe_str(report.get("literature_summary", "")))

    doc.add_heading("Clinical Trials Summary", level=2)
    doc.add_paragraph(_safe_str(report.get("clinical_trials_summary", "")))

    doc.add_heading("Competition Summary", level=2)
    doc.add_paragraph(_safe_str(report.get("competition_summary", "")))

    doc.add_heading("Major Risks", level=2)
    risks = report.get("major_risks", [])
    if not isinstance(risks, list):
        risks = [str(risks)] if risks else []
    for risk in risks:
        doc.add_paragraph(_safe_str(risk), style="List Bullet")

    doc.add_heading("Major Opportunities", level=2)
    opps = report.get("major_opportunities", [])
    if not isinstance(opps, list):
        opps = [str(opps)] if opps else []
    for opp in opps:
        doc.add_paragraph(_safe_str(opp), style="List Bullet")

    doc.add_heading(f"Recommendation: {report.get('recommendation', '')}", level=1)
    doc.add_paragraph(f"Reasoning: {report.get('reasoning', '')}")
    doc.add_paragraph(f"Uncertainty: {report.get('uncertainty', '')}")

    doc.add_heading("Citations", level=2)
    for c in report.get("citations", []):
        doc.add_paragraph(_normalize_citation(c))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _setup_pdf_fonts(pdf: FPDF) -> str:
    """Register fonts and return the font family name to use."""
    if _HAS_CJK_FONT:
        pdf.add_font("NotoSansSC", "", _FONT_REGULAR)
        if os.path.isfile(_FONT_BOLD):
            pdf.add_font("NotoSansSC", "B", _FONT_BOLD)
        else:
            # Use regular as fallback for bold
            pdf.add_font("NotoSansSC", "B", _FONT_REGULAR)
        return "NotoSansSC"
    else:
        logger.warning(
            "CJK font not found at %s — PDF will use Helvetica (CJK characters will be missing). "
            "Download NotoSansSC from Google Fonts and place in %s",
            _FONT_REGULAR,
            _FONT_DIR,
        )
        return "Helvetica"


def generate_pdf_report(report: dict) -> bytes:
    """Generate a PDF report from the decision output."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    font_family = _setup_pdf_fonts(pdf)

    pdf.set_font(font_family, "B", 16)
    pdf.cell(0, 10, "Drug Target Assessment Report", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font(font_family, "B", 12)
    pdf.cell(0, 8, f"Target: {report.get('target', '')}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, f"Indication: {report.get('indication', '')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    def _section(title: str, body: str):
        pdf.set_font(font_family, "B", 11)
        pdf.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_family, "", 10)
        pdf.multi_cell(0, 6, body or "N/A")
        pdf.ln(2)

    _section("Literature Summary", _safe_str(report.get("literature_summary", "")))
    _section("Clinical Trials Summary", _safe_str(report.get("clinical_trials_summary", "")))
    _section("Competition Summary", _safe_str(report.get("competition_summary", "")))

    # Risks & Opportunities
    pdf_risks = report.get("major_risks", [])
    if not isinstance(pdf_risks, list):
        pdf_risks = [str(pdf_risks)] if pdf_risks else []
    _section("Major Risks", "\n".join(f"- {_safe_str(r)}" for r in pdf_risks) or "N/A")
    pdf_opps = report.get("major_opportunities", [])
    if not isinstance(pdf_opps, list):
        pdf_opps = [str(pdf_opps)] if pdf_opps else []
    _section("Major Opportunities", "\n".join(f"- {_safe_str(o)}" for o in pdf_opps) or "N/A")

    # Recommendation
    pdf.set_font(font_family, "B", 12)
    pdf.cell(0, 8, f"Recommendation: {report.get('recommendation', '')}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font_family, "", 10)
    pdf.multi_cell(0, 6, f"Reasoning: {report.get('reasoning', '')}")
    pdf.ln(2)
    pdf.multi_cell(0, 6, f"Uncertainty: {report.get('uncertainty', '')}")
    pdf.ln(4)

    # Citations
    citations = report.get("citations", [])
    if citations:
        pdf.set_font(font_family, "B", 11)
        pdf.cell(0, 8, "Citations", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_family, "", 9)
        for c in citations:
            pdf.multi_cell(0, 5, f"- {_normalize_citation(c)}")

    return bytes(pdf.output())
